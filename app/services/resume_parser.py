"""
app/services/resume_parser.py — Text extraction utilities for PDF and DOCX files.
"""

import io
from pypdf import PdfReader
import docx


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract clean raw text from PDF bytes using pypdf."""
    reader = PdfReader(io.BytesIO(file_bytes))
    text_chunks = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_chunks.append(page_text.strip())
    return "\n\n".join(text_chunks)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract clean raw text from DOCX bytes using python-docx."""
    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    # Also extract any table content
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                paragraphs.append(" | ".join(row_text))
                
    return "\n\n".join(paragraphs)


def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Dispatcher based on filename extension.
    Raises ValueError if unsupported file extension.
    """
    lower_name = filename.lower()
    if lower_name.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif lower_name.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file format for resume: '{filename}'. Please upload a .pdf or .docx file.")
