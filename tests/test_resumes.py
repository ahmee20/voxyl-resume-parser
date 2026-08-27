"""
tests/test_resumes.py — Unit and integration tests for resume upload & extraction.
"""

import io
import pytest
from httpx import AsyncClient
import docx
from pypdf import PdfWriter

from app.services.resume_parser import extract_text, extract_text_from_docx, extract_text_from_pdf
from app.agent.nodes.extract_resume import extract_resume_node, generate_resume_html_skeleton
from app.agent.state import GraphState
from app.models.user import User


def create_sample_docx(text_content: str) -> bytes:
    """Helper to generate a valid in-memory DOCX file."""
    doc = docx.Document()
    doc.add_heading("John Doe", level=1)
    for paragraph in text_content.split("\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def create_sample_pdf(text_content: str) -> bytes:
    """Helper to generate a minimal valid in-memory PDF file."""
    writer = PdfWriter()
    # Add a blank page
    writer.add_blank_page(width=72 * 8.5, height=72 * 11)
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


@pytest.mark.asyncio
async def test_extract_text_from_docx():
    content = "Software Engineer\nPython, FastAPI, LangGraph"
    docx_bytes = create_sample_docx(content)
    extracted = extract_text_from_docx(docx_bytes)
    assert "John Doe" in extracted
    assert "Software Engineer" in extracted
    assert "Python, FastAPI, LangGraph" in extracted


@pytest.mark.asyncio
async def test_extract_text_unsupported_format():
    with pytest.raises(ValueError, match="Unsupported file format"):
        extract_text(b"some text", "resume.txt")


@pytest.mark.asyncio
async def test_extract_resume_node_with_mocked_llm():
    """Test LangGraph node generates structured resume HTML."""
    state: GraphState = {
        "user_id": 1,
        "application_id": 10,
        "resume_text": "Alex Smith\nAI Engineer\nSkills: Python, PyTorch",
    }

    updated_state = extract_resume_node(state)
    assert 'class="resume-document"' in updated_state["resume_html"]
    assert "Alex Smith" in updated_state["resume_html"]
    assert updated_state["resume_text"] == state["resume_text"]


@pytest.mark.asyncio
async def test_upload_resume_unsupported_format(client: AsyncClient):
    files = {"file": ("test.txt", b"plain text", "text/plain")}
    response = await client.post("/resumes/upload?user_id=1", files=files)
    assert response.status_code == 400
    assert "Only .pdf and .docx" in response.json()["detail"]


@pytest.mark.asyncio
async def test_upload_resume_success(client: AsyncClient, db_session):
    # Seed a user first
    user = User(
        google_sub="test-user-sub-123",
        email="testuser@example.com",
        name="Test User",
    )
    db_session.add(user)
    await db_session.commit()
    await db_session.refresh(user)

    docx_bytes = create_sample_docx("Experienced Developer\nPython, Docker, SQL")
    files = {"file": ("my_resume.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}

    response = await client.post(f"/resumes/upload?user_id={user.id}", files=files)
    assert response.status_code == 201
    data = response.json()
    assert data["user_id"] == user.id
    assert data["version"] == 1
    assert data["is_base"] is True
    assert "John Doe" in data["source_text"]
    assert 'class="resume-document"' in data["source_html"]

    resume_id = data["id"]
    # Test GET /resumes/{resume_id}
    get_res = await client.get(f"/resumes/{resume_id}")
    assert get_res.status_code == 200
    assert get_res.json()["id"] == resume_id
